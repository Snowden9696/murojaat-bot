import logging
from datetime import datetime
from aiogram import Bot, Dispatcher, types
from aiogram.filters import Command
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.context import FSMContext
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove, FSInputFile
from openpyxl import Workbook, load_workbook
from docx import Document
import os
import asyncio

API_TOKEN = os.getenv("API_TOKEN")
ADMIN_IDS = [900437855, 630119080]

bot = Bot(token=API_TOKEN)
dp = Dispatcher(storage=MemoryStorage())
logging.basicConfig(level=logging.INFO)

class Form(StatesGroup):
    full_name = State()
    birth_date = State()
    address = State()
    phone = State()
    topic = State()
    short_reason = State()
    full_text = State()

topics = [
    "Huquqiy himoya",
    "Ijtimoiy iqtisodiy manfaatlarni himoya qilish",
    "Mehnat muhofazasi",
    "Kasaba uyushma tashkiliy-huquqiy faoliyati",
    "Sog'lomlashtirish",
    "Moddiy yordam ajratish",
    "Ijtimoiy-mehnat munosabatlariga oid boshqa masalalar",
    "Kasaba uyushmasi vakolatiga oid bo'lmagan masalalar",
    "Paxta masalasi"
]
topic_keyboard = ReplyKeyboardMarkup(
    keyboard=[[KeyboardButton(text=topic)] for topic in topics],
    resize_keyboard=True,
    one_time_keyboard=True
)

def calculate_age(birth_str):
    try:
        birth_date = datetime.strptime(birth_str, "%d.%m.%Y")
        today = datetime.today()
        return today.year - birth_date.year - ((today.month, today.day) < (birth_date.month, birth_date.day))
    except:
        return None

def write_to_excel(data, user_id):
    age = calculate_age(data["birth_date"])
    file_path = "applications_31plus.xlsx" if age and age >= 31 else "applications_upto30.xlsx"
    if not os.path.exists(file_path):
        wb = Workbook()
        ws = wb.active
        ws.append(["Ariza ID", "Telegram ID", "Yosh", "F.I.Sh", "Tug‘ilgan sana", "Manzil", "Telefon", "Yo‘nalish", "Mazmuni", "To‘liq matn", "Yuborilgan sana/vaqt"])
        next_id = 1
    else:
        wb = load_workbook(file_path)
        ws = wb.active
        next_id = ws.max_row
    time_str = datetime.now().strftime("%d.%m.%Y %H:%M")
    ws.append([next_id, user_id, age if age else "Aniqlanmadi", data["full_name"], data["birth_date"], data["address"], data["phone"], data["topic"], data["short_reason"], data["full_text"], time_str])
    wb.save(file_path)
    return next_id, age, time_str

def create_docx(ariza_id, data, user_id, age, timestamp):
    doc = Document()
    doc.add_heading(f"Ariza №{ariza_id}", level=1)
    doc.add_paragraph(f"🕓 Yaratilgan sana: {timestamp}\n👤 Telegram ID: {user_id}\nYosh: {age if age else 'Aniqlanmadi'}")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Savol'
    hdr_cells[1].text = 'Javob'

    qna = {
        "1. Familiyangiz, ismingiz va sharifingiz:": data["full_name"],
        "2. Tug‘ilgan sana:": data["birth_date"],
        "3. Yashash manzili:": data["address"],
        "4. Telefon raqami:": data["phone"],
        "5. Yo‘nalish:": data["topic"],
        "6. Murojaat mazmuni:": data["short_reason"],
        "7. Murojaat matni:": data["full_text"]
    }

    for question, answer in qna.items():
        row_cells = table.add_row().cells
        row_cells[0].text = question
        row_cells[1].text = answer

    filename = f"ariza_{ariza_id}.docx"
    doc.save(filename)
    return filename

@dp.message(Command("start"))
async def start_handler(message: types.Message, state: FSMContext):
    text = (
        "🇺🇿 *O'zbekiston kasaba uyushmalari Federatsiyasi Toshkent viloyati Kengashi*\n\n"
        "🤖 Murojaatlarni qabul qilish botiga xush kelibsiz!\n\n"
        "Marhamat, \"Murojaat yuborish\" tugmasini bosib murojaatingizni yo'llashingiz mumkin."
    )
    markup = ReplyKeyboardMarkup(keyboard=[[KeyboardButton(text="Murojaat yuborish")]], resize_keyboard=True)
    await message.answer(text, reply_markup=markup, parse_mode="Markdown")

@dp.message(lambda m: m.text == "Murojaat yuborish")
async def trigger_murojaat(message: types.Message, state: FSMContext):
    await message.answer("1. Familiyangiz, ismingiz va sharifingiz:\n(Masalan, Azizov Aziz Azizovich)", reply_markup=ReplyKeyboardRemove())
    await state.set_state(Form.full_name)

@dp.message(Form.full_name)
async def step_full_name(message: types.Message, state: FSMContext):
    await state.update_data(full_name=message.text)
    await message.answer("2. Tug‘ilgan sana:\n(Masalan, 01.01.1991)")
    await state.set_state(Form.birth_date)

@dp.message(Form.birth_date)
async def step_birth(message: types.Message, state: FSMContext):
    await state.update_data(birth_date=message.text)
    await message.answer("3. Yashash manzili:\n(Masalan, Toshkent viloyati, Nurafshon shahri...)")
    await state.set_state(Form.address)

@dp.message(Form.address)
async def step_address(message: types.Message, state: FSMContext):
    await state.update_data(address=message.text)
    await message.answer("4. Telefon raqamingiz:\n(Masalan, 90-123-45-67)")
    await state.set_state(Form.phone)

@dp.message(Form.phone)
async def step_phone(message: types.Message, state: FSMContext):
    await state.update_data(phone=message.text)
    await message.answer("5. Yo‘nalishni tanlang:", reply_markup=topic_keyboard)
    await state.set_state(Form.topic)

@dp.message(Form.topic)
async def step_topic(message: types.Message, state: FSMContext):
    await state.update_data(topic=message.text)
    await message.answer("6. Murojaat mazmuni:\n(Masalan, ishga tiklash...)")
    await state.set_state(Form.short_reason)

@dp.message(Form.short_reason)
async def step_reason(message: types.Message, state: FSMContext):
    await state.update_data(short_reason=message.text)
    await message.answer("7. Murojaat matni:")
    await state.set_state(Form.full_text)

@dp.message(Form.full_text)
async def step_full_text(message: types.Message, state: FSMContext):
    data = await state.get_data()
    data["full_text"] = message.text

    ariza_id, age, timestamp = write_to_excel(data, message.from_user.id)
    doc_file = create_docx(ariza_id, data, message.from_user.id, age, timestamp)

    for admin in ADMIN_IDS:
        await bot.send_message(admin, f"📩 Yangi murojaat qabul qilindi:\n\n👤 {data['full_name']}\n📞 {data['phone']}")
        await bot.send_document(admin, FSInputFile(doc_file), caption=f"📄 Ariza #{ariza_id}")

    os.remove(doc_file)
    await message.answer("✅ Murojaatingiz qabul qilindi, tez orada sizga aloqaga chiqamiz!")
    await state.clear()

async def main():
    await dp.start_polling(bot)

if __name__ == "__main__":
    asyncio.run(main())